#!/usr/bin/env python3
"""
Build a clean .docx for the AI Reel Pipeline architecture.
Combines docs/ARCHITECTURE.md, docs/DIAGRAMS.md, and docs/PROMPT_TEMPLATE.md.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "AI-Reel-Pipeline-Architecture.docx"
DIAGRAMS = ROOT / "docs" / "diagrams"


def insert_diagram(doc, png_name, width_inches=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run()
    r.add_picture(str(DIAGRAMS / png_name), width=Inches(width_inches))

doc = Document()

# ── Default font ──────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def H1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

def H2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x2C, 0x5A, 0x8C)

def H3(text):
    p = doc.add_heading(text, level=3)

def P(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def diagram_block(text):
    """ASCII diagram as monospace paragraph."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Menlo"
    r.font.size = Pt(8)

# ─────────────────────────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
tr = title.add_run("\n\n\nAI Reel Pipeline")
tr.font.size = Pt(36); tr.bold = True
tr.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

sub = doc.add_paragraph()
sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
sr = sub.add_run("Architecture & Operations")
sr.font.size = Pt(20)
sr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

tag = doc.add_paragraph()
tag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
tagr = tag.add_run("\n\nA productized, end-to-end system that converts a single content brief "
                   "into a publish-ready vertical video reel.\n")
tagr.font.size = Pt(12)
tagr.italic = True

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 1. WHAT IT DOES
# ─────────────────────────────────────────────────────────────────
H1("1. What This System Does")
P("Input: a short topic, transcript, or content brief — driven by a client-supplied prompt.")
P("Output: a 60 to 90 second vertical (9:16) video reel with talking-head avatar, "
  "motion-graphic b-roll, burned-in captions, ambient audio, thumbnail card, and "
  "ready-to-post social copy.")
P("")
P("The pipeline is:")
bullet("Deterministic — same input + same prompt = same output bundle")
bullet("Resumable — every stage writes state; auto-resumes on retry")
bullet("Stage-isolated — clean separation of script, media, composition")
bullet("Parallel where possible — three media stages run concurrently")
bullet("Re-renderable in minutes — visual changes don't trigger a full re-run")

# ─────────────────────────────────────────────────────────────────
# 2. FUNCTIONAL ARCHITECTURE
# ─────────────────────────────────────────────────────────────────
doc.add_page_break()
H1("2. Functional Architecture")
P("What the system does, viewed as a flow of content artifacts. Every box is an "
  "isolated unit. Every arrow is a data hand-off via a versioned JSON contract.")

insert_diagram(doc, "functional_architecture.png", width_inches=6.5)

H2("Stages")
P("")
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "#"
hdr[1].text = "Stage"
hdr[2].text = "What It Owns"
for n, name, owns in [
    ("1", "Script Generation",  "Topic / brief → structured 10-section script (self-scored)"),
    ("2", "B-Roll Generation",  "Per-section motion graphic, screen, diagram, terminal clips"),
    ("3", "Avatar Generation",  "Synthetic talking-head video from the spoken script"),
    ("4", "Audio Generation",   "Ambient bed tracks + impact stings at trigger moments"),
    ("5", "Compose & Render",   "Transcript alignment + 1080x1920 reel + 1.25x fast cut"),
    ("6", "Quality Review",     "Multi-image visual evaluation against design rubric"),
    ("8", "Social Copy",        "YouTube / Instagram / LinkedIn captions and hashtags"),
]:
    row = table.add_row().cells
    row[0].text = n; row[1].text = name; row[2].text = owns

# ─────────────────────────────────────────────────────────────────
# 3. TECHNICAL ARCHITECTURE
# ─────────────────────────────────────────────────────────────────
doc.add_page_break()
H1("3. Technical Architecture")
P("How the system is built, viewed as runtime components and external services. "
  "Component boundaries are firm; external services are interchangeable.")

insert_diagram(doc, "technical_architecture.png", width_inches=6.5)

H2("External Services")
P("Each service has at least one drop-in alternative.")
t2 = doc.add_table(rows=1, cols=3)
t2.style = "Light Grid Accent 1"
h2 = t2.rows[0].cells
h2[0].text = "Service"; h2[1].text = "Purpose"; h2[2].text = "Replaceable"
for s, p, r in [
    ("Anthropic API",        "Script, scoring, social copy, visual review", "No (model-specific)"),
    ("HeyGen",               "Synthetic talking-head avatar",               "Yes (any avatar API)"),
    ("Kie.ai (ElevenLabs)",  "Ambient music + impact stings",               "Yes"),
    ("Kie.ai (Imagen)",      "Reference imagery (optional)",                "Yes"),
    ("OpenAI Whisper",       "Word-level transcript alignment",             "Yes (runs locally)"),
]:
    row = t2.add_row().cells; row[0].text = s; row[1].text = p; row[2].text = r

H2("Internal Components")
t3 = doc.add_table(rows=1, cols=3)
t3.style = "Light Grid Accent 1"
h3 = t3.rows[0].cells
h3[0].text = "Component"; h3[1].text = "Technology"; h3[2].text = "Role"
for c, t, r in [
    ("Orchestrator",     "Python",                "Stage runner, state machine"),
    ("Animation engine", "HyperFrames + GSAP",    "Motion graphics → MP4"),
    ("Composition",      "Remotion + React",      "Layered timeline → MP4"),
    ("Transcode",        "FFmpeg",                "Codec normalization, speed-up"),
    ("Asset store",      "Versioned filesystem",  "Stage-to-stage hand-off"),
]:
    row = t3.add_row().cells; row[0].text = c; row[1].text = t; row[2].text = r

# ─────────────────────────────────────────────────────────────────
# 4. TIMING
# ─────────────────────────────────────────────────────────────────
doc.add_page_break()
H1("4. Operational Model — Timing & Cost")
P("Per-reel wall-clock and cost footprint. End-to-end: 15 to 25 minutes per reel "
  "on standard developer hardware.")
P("")
insert_diagram(doc, "timing_gantt.png", width_inches=6.5)
P("")

t4 = doc.add_table(rows=1, cols=4)
t4.style = "Light Grid Accent 1"
h4 = t4.rows[0].cells
h4[0].text = "Stage"; h4[1].text = "Typical Duration"; h4[2].text = "Parallel"; h4[3].text = "Cost Driver"
for s, d, p, c in [
    ("Script",         "~90 seconds",    "No",  "LLM tokens"),
    ("B-Roll",         "~3 minutes",     "Yes", "Browser render, optional media APIs"),
    ("Avatar",         "~5–10 minutes",  "Yes", "Avatar generation credits"),
    ("Audio",          "~2 minutes",     "Yes", "TTS / SFX credits"),
    ("Compose",        "~5–10 minutes",  "No",  "Local CPU"),
    ("Quality Review", "~1 minute",      "No",  "LLM tokens (vision)"),
    ("Social Copy",    "~20 seconds",    "No",  "LLM tokens"),
]:
    row = t4.add_row().cells; row[0].text = s; row[1].text = d; row[2].text = p; row[3].text = c

# ─────────────────────────────────────────────────────────────────
# 5. QUALITY CONTROLS
# ─────────────────────────────────────────────────────────────────
H1("5. Quality Controls")
bullet("Script gate — generated scripts are scored against a structured rubric and "
       "re-generated until they clear a quality threshold.")
bullet("Pre-render check — script is sanity-checked before paying for avatar or "
       "media generation.")
bullet("Visual review — finished reels are evaluated against a multi-dimensional "
       "rubric (hook strength, on-screen text legibility, audio sync, CTA clarity).")
bullet("Stale-asset guard — downstream assets are auto-invalidated when their "
       "upstream contract changes.")

# ─────────────────────────────────────────────────────────────────
# 6. EXTENSIBILITY
# ─────────────────────────────────────────────────────────────────
H1("6. Extensibility Surfaces")
P("The pipeline is intentionally modular. Four high-leverage extension points:")
bullet("Prompt template — swap scripting style without touching code.")
bullet("Visual template library — add new motion-graphic patterns by dropping HTML files.")
bullet("Avatar provider — replace via configuration.")
bullet("Composition layer — change layout, captions, rhythm via component edits "
       "without affecting upstream stages.")
P("")
insert_diagram(doc, "extensibility.png", width_inches=6.5)

# ─────────────────────────────────────────────────────────────────
# 7. WHERE THE CLIENT PROMPT GOES
# ─────────────────────────────────────────────────────────────────
doc.add_page_break()
H1("7. Where The Client Prompt Goes")
P("The pipeline accepts a client-supplied prompt template that fully replaces "
  "the default Stage 1 generator. The template is the only piece of business "
  "logic the client owns. Every downstream stage is template-agnostic.")
P("")
H2("Required Output Schema")
P("Whatever your prompt produces, the model output must match this JSON shape:")
code("""{
  "title":               "<60 chars max>",
  "total_duration_sec":  60–95,
  "bgm_transition_sec":  <number>,
  "bgm_dip_timestamps":  [3 numbers],
  "full_spoken_script":  "<complete narration>",
  "grand_takeaway_line": "<quotable line>",
  "tool_mentioned":      "<product / methodology>",
  "sections": [ ...10 section objects... ]
}""")

H2("Prompt Slot")
P("Paste the client's scripting prompt here. The pipeline injects the input "
  "topic / brief into the prompt using {{topic}} or {{brief}} placeholders and "
  "calls the language model.")
P("")
P("┌──────────────────────────────────────────────────────────────┐", italic=True)
P("│   [ CLIENT PROMPT — PASTE HERE ]                              │", italic=True)
P("│                                                                │", italic=True)
P("│   The client owns the wording, tone, audience profile, brand   │", italic=True)
P("│   voice, CTA style, and structural constraints. The pipeline   │", italic=True)
P("│   only enforces the output schema.                             │", italic=True)
P("│                                                                │", italic=True)
P("│   Refer to docs/PROMPT_TEMPLATE.md for the canonical slot      │", italic=True)
P("│   used at runtime.                                             │", italic=True)
P("└──────────────────────────────────────────────────────────────┘", italic=True)

# ─────────────────────────────────────────────────────────────────
# 8. WHAT THIS DOES NOT DO
# ─────────────────────────────────────────────────────────────────
H1("8. What This System Does Not Do")
P("To set expectations crisply:")
bullet("Generate scripts longer than ~90 seconds of spoken content")
bullet("Mix multiple speakers or styles in a single reel")
bullet("Render landscape (16:9) outputs — vertical 9:16 only")
bullet("Auto-publish to social platforms — outputs are files; distribution is external")
bullet("Process live or streaming inputs — batch only")

# ─────────────────────────────────────────────────────────────────
# OUTPUT
# ─────────────────────────────────────────────────────────────────
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"Wrote: {OUT}")
print(f"Size: {OUT.stat().st_size // 1024} KB")
